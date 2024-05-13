from langchain_community.llms import CTransformers
from langchain.chains import LLMChain
from langchain_core.prompts import PromptTemplate
import streamlit as st 
from docx import Document
from docx.shared import Inches
from PIL import Image
import requests
import io


#Loading the model
def load_llm(max_tokens, prompt_template):
     # Load the locally downloaded model here
    llm = CTransformers(
        model = "llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens = max_tokens,
        temperature = 0.92
    )

    llm_chain = LLMChain(
         llm=llm,
         prompt=PromptTemplate.from_template(prompt_template)
    )
    print(llm_chain)
    return llm_chain

def get_src_original_url(query):
    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': "z0Po9yA3nTt1gdKfCXGnFYAVakKNMpdsLcuBWKbnDrqmfMzqstanMpqT",
    }

    params = {
        'query': query,
    'per_page': 1,
    }

    response = requests.get(url, headers=headers, params=params)

    # Check if the request was successful (status code 200)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            src_original_url = photos[0]['src']['original']
            return src_original_url
        
        else:
            st.write("No photos found for the given query.")
    else:
        st.write(f"Error: {response.status_code}, {response.text}")

    return None

def create_word_docx(user_input, paragraph, image_input):
    # Create a new Word document
    doc = Document()

    # Add the user input to the document
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)

    # Add the image to the document
    doc.add_heading('Image Output', level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))
    return doc

st.set_page_config(layout="wide")

def main():
    st.title("AcroMemory")

    user_input = st.text_input("Please Enter the phrase for which you want to generate an Acronym")

    image_input = st.text_input("Enter the topic the phrase is related to")
    
    submit=st.button("Generate")
    
    if submit:
        if len(user_input) > 0 and len(image_input) > 0:

            col1, col2, col3 = st.columns([1,2,1])

            with col1:
                st.subheader("Generated Content by Llama 2")
                prompt_template = """Generate a random fictional or factual sentence using the first letters of the words from the input text "{user_input}" thereby making the inital text easy to remember and try not to use the words from the input text in the generated sentence. Display only the generated sentence.
                """
                llm_call = load_llm(max_tokens=1000, prompt_template=prompt_template)
                print(llm_call)
                result = llm_call(user_input)
                if len(result) > 0:
                    st.info("Your acronym has been been generated successfully!")
                    st.write(result)
                else:
                    st.error("Your acronym couldn't be generated!")

            with col2:
                st.subheader("Fetched Image")
                image_url = get_src_original_url(image_input)
                st.image(image_url)

            with col3:
                st.subheader("Final Doc to Download")
                image_input = Image.open(io.BytesIO(requests.get(image_url).content))           
                doc = create_word_docx(user_input, result['text'], image_input)

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )


if __name__ == "__main__":
    main()